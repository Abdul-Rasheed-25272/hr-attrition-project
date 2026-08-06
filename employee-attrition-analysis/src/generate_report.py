from datetime import date
from html import escape
from pathlib import Path
import json

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# --------------------------------------------------
# Project paths
# --------------------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parents[1]

CLEAN_DATA_FILE = (
    PROJECT_ROOT
    / "data"
    / "processed"
    / "hr_employee_attrition_clean.csv"
)

EDA_SUMMARY_FILE = (
    PROJECT_ROOT
    / "reports"
    / "eda_summary.json"
)

QUALITY_REPORT_FILE = (
    PROJECT_ROOT
    / "reports"
    / "data_quality_report.json"
)

DEPARTMENT_FILE = (
    PROJECT_ROOT
    / "reports"
    / "attrition_by_department.csv"
)

JOB_ROLE_FILE = (
    PROJECT_ROOT
    / "reports"
    / "attrition_by_job_role.csv"
)

FIGURES_DIRECTORY = (
    PROJECT_ROOT
    / "reports"
    / "figures"
)

FINAL_DIRECTORY = (
    PROJECT_ROOT
    / "reports"
    / "final"
)

WORD_REPORT = (
    FINAL_DIRECTORY
    / "HR_Attrition_Analysis_Report.docx"
)

HTML_REPORT = (
    FINAL_DIRECTORY
    / "HR_Attrition_Analysis_Report.html"
)


# --------------------------------------------------
# Dataset descriptions
# --------------------------------------------------

DATASET_GROUPS = [
    {
        "category": "Target variable",
        "variables": "attrition, attrition_flag",
        "description": (
            "Shows whether an employee left the company. "
            "The numerical flag uses 1 for Yes and 0 for No."
        )
    },
    {
        "category": "Employee identification",
        "variables": "employee_number",
        "description": (
            "Unique employee identifier. It should not be used "
            "as a predictive machine-learning feature."
        )
    },
    {
        "category": "Personal characteristics",
        "variables": (
            "age, gender, marital_status, education, "
            "education_field, distance_from_home"
        ),
        "description": (
            "Basic demographic, education and commuting information."
        )
    },
    {
        "category": "Employment information",
        "variables": (
            "department, job_role, job_level, business_travel, "
            "job_involvement, over_time"
        ),
        "description": (
            "Describes the employee's position and work conditions."
        )
    },
    {
        "category": "Salary and compensation",
        "variables": (
            "monthly_income, daily_rate, hourly_rate, monthly_rate, "
            "percent_salary_hike, stock_option_level"
        ),
        "description": (
            "Contains income, payment-rate and compensation information."
        )
    },
    {
        "category": "Satisfaction and workplace",
        "variables": (
            "environment_satisfaction, job_satisfaction, "
            "relationship_satisfaction, work_life_balance"
        ),
        "description": (
            "Numerical measures related to satisfaction and workplace experience."
        )
    },
    {
        "category": "Experience and tenure",
        "variables": (
            "total_working_years, years_at_company, "
            "years_in_current_role, years_since_last_promotion, "
            "years_with_curr_manager, num_companies_worked"
        ),
        "description": (
            "Describes career experience, company tenure, promotion "
            "history and manager-related tenure."
        )
    },
    {
        "category": "Training and performance",
        "variables": (
            "training_times_last_year, performance_rating"
        ),
        "description": (
            "Contains training frequency and employee performance information."
        )
    }
]


CHARTS = [
    (
        "attrition_distribution.png",
        "Figure 1: Employee Attrition Distribution",
        "Comparison of employees who stayed and employees who left."
    ),
    (
        "attrition_by_department.png",
        "Figure 2: Attrition Rate by Department",
        "Percentage of employees leaving within each department."
    ),
    (
        "attrition_by_job_role.png",
        "Figure 3: Attrition Rate by Job Role",
        "Comparison of attrition rates across different job roles."
    ),
    (
        "attrition_by_overtime.png",
        "Figure 4: Attrition Rate by Overtime",
        "Comparison between employees who work overtime and those who do not."
    ),
    (
        "income_by_attrition.png",
        "Figure 5: Monthly Income by Attrition Status",
        "Distribution of monthly income for employees who stayed and left."
    )
]


# --------------------------------------------------
# Loading and validation
# --------------------------------------------------

def validate_files() -> None:
    """Confirm that all required analysis files exist."""

    required_files = [
        CLEAN_DATA_FILE,
        EDA_SUMMARY_FILE,
        QUALITY_REPORT_FILE,
        DEPARTMENT_FILE,
        JOB_ROLE_FILE
    ]

    missing_files = [
        str(file_path)
        for file_path in required_files
        if not file_path.exists()
    ]

    if missing_files:
        missing_text = "\n".join(missing_files)

        raise FileNotFoundError(
            "The following analysis files are missing:\n"
            f"{missing_text}\n\n"
            "Run clean_data.py and explore_data.py first."
        )


def load_results():
    """Load all generated analysis results."""

    cleaned_df = pd.read_csv(CLEAN_DATA_FILE)

    with EDA_SUMMARY_FILE.open(
        "r",
        encoding="utf-8"
    ) as file:
        eda_summary = json.load(file)

    with QUALITY_REPORT_FILE.open(
        "r",
        encoding="utf-8"
    ) as file:
        quality_report = json.load(file)

    department_df = pd.read_csv(DEPARTMENT_FILE)
    job_role_df = pd.read_csv(JOB_ROLE_FILE)

    return (
        cleaned_df,
        eda_summary,
        quality_report,
        department_df,
        job_role_df
    )


# --------------------------------------------------
# Analysis interpretation
# --------------------------------------------------

def build_key_findings(
    cleaned_df: pd.DataFrame,
    eda_summary: dict,
    department_df: pd.DataFrame,
    job_role_df: pd.DataFrame
) -> list[str]:
    """Generate readable findings from the analysis results."""

    findings = []

    attrition_rate = eda_summary[
        "overall_attrition_rate_percent"
    ]

    findings.append(
        f"The overall employee attrition rate is "
        f"{attrition_rate:.2f}%."
    )

    highest_department = department_df.loc[
        department_df["attrition_rate"].idxmax()
    ]

    findings.append(
        f"{highest_department['department']} has the highest "
        f"department-level attrition rate at "
        f"{highest_department['attrition_rate']:.2f}%."
    )

    highest_role = job_role_df.loc[
        job_role_df["attrition_rate"].idxmax()
    ]

    findings.append(
        f"{highest_role['job_role']} has the highest job-role "
        f"attrition rate at {highest_role['attrition_rate']:.2f}%."
    )

    overtime_rates = (
        cleaned_df.groupby("over_time")["attrition_flag"]
        .mean()
        .mul(100)
    )

    if "Yes" in overtime_rates.index and "No" in overtime_rates.index:
        findings.append(
            f"Employees working overtime have an attrition rate of "
            f"{overtime_rates['Yes']:.2f}%, compared with "
            f"{overtime_rates['No']:.2f}% for employees not "
            f"working overtime."
        )

    median_income = (
        cleaned_df.groupby("attrition")["monthly_income"]
        .median()
    )

    if "Yes" in median_income.index and "No" in median_income.index:
        findings.append(
            f"The median monthly income is "
            f"{median_income['Yes']:,.0f} for employees who left "
            f"and {median_income['No']:,.0f} for employees who stayed."
        )

    average_tenure = (
        cleaned_df.groupby("attrition")["years_at_company"]
        .mean()
    )

    if "Yes" in average_tenure.index and "No" in average_tenure.index:
        findings.append(
            f"Employees who left had spent an average of "
            f"{average_tenure['Yes']:.2f} years at the company, "
            f"compared with {average_tenure['No']:.2f} years for "
            f"employees who stayed."
        )

    findings.append(
        "These findings show associations in the dataset and should "
        "not be interpreted as proof that one factor directly causes attrition."
    )

    return findings


# --------------------------------------------------
# Word formatting helpers
# --------------------------------------------------

def shade_cell(cell, fill: str) -> None:
    """Apply a background colour to a Word table cell."""

    properties = cell._tc.get_or_add_tcPr()

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)

    properties.append(shading)


def format_word_table(
    table,
    header_colour: str = "D9EAF7"
) -> None:
    """Apply consistent table formatting."""

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for cell in table.rows[0].cells:
        shade_cell(cell, header_colour)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)

                for run in paragraph.runs:
                    run.font.size = Pt(9)


def add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    columns: list[str],
    column_labels: list[str],
    percentage_columns: list[str] | None = None
) -> None:
    """Add a pandas dataframe to a Word document."""

    percentage_columns = percentage_columns or []

    table = document.add_table(
        rows=1,
        cols=len(columns)
    )

    for index, label in enumerate(column_labels):
        table.rows[0].cells[index].text = label

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells

        for index, column in enumerate(columns):
            value = row[column]

            if column in percentage_columns:
                text = f"{float(value):.2f}%"
            elif isinstance(value, float):
                text = f"{value:,.2f}"
            else:
                text = str(value)

            cells[index].text = text

    format_word_table(table)


def add_metric_table(
    document: Document,
    metrics: list[tuple[str, str]]
) -> None:
    """Add a key-metrics table to the Word report."""

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Result"

    for label, value in metrics:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    format_word_table(
        table,
        header_colour="CFE8DC"
    )


# --------------------------------------------------
# Word report
# --------------------------------------------------

def create_word_report(
    cleaned_df: pd.DataFrame,
    eda_summary: dict,
    quality_report: dict,
    department_df: pd.DataFrame,
    job_role_df: pd.DataFrame,
    findings: list[str]
) -> None:
    """Generate the formatted Word report."""

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)

    for style_name in [
        "Title",
        "Heading 1",
        "Heading 2"
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"

    document.styles["Heading 1"].font.color.rgb = RGBColor(
        31,
        78,
        121
    )

    document.styles["Heading 2"].font.color.rgb = RGBColor(
        46,
        116,
        181
    )

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        "HR Employee Attrition Analysis"
    )

    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(
        31,
        78,
        121
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "Data Cleaning, Exploratory Analysis and Key Findings"
    )

    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)

    report_date = document.add_paragraph()
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_date.add_run(
        f"Generated on {date.today().strftime('%d %B %Y')}"
    )

    document.add_page_break()

    # Executive summary
    document.add_heading(
        "1. Executive Summary",
        level=1
    )

    document.add_paragraph(
        "This report examines employee attrition using demographic, "
        "employment, compensation, satisfaction and experience-related "
        "variables. The purpose is to identify patterns associated with "
        "employees leaving the organisation."
    )

    metrics = [
        (
            "Total employees",
            f"{eda_summary['total_employees']:,}"
        ),
        (
            "Employees who left",
            f"{eda_summary['employees_left']:,}"
        ),
        (
            "Employees who stayed",
            f"{eda_summary['employees_stayed']:,}"
        ),
        (
            "Overall attrition rate",
            (
                f"{eda_summary['overall_attrition_rate_percent']:.2f}%"
            )
        ),
        (
            "Average employee age",
            f"{eda_summary['average_age']:.2f} years"
        ),
        (
            "Average monthly income",
            f"{eda_summary['average_monthly_income']:,.2f}"
        )
    ]

    add_metric_table(
        document,
        metrics
    )

    # Dataset overview
    document.add_heading(
        "2. Dataset Overview",
        level=1
    )

    document.add_paragraph(
        f"The cleaned dataset contains {cleaned_df.shape[0]:,} employee "
        f"records and {cleaned_df.shape[1]} variables. Each row represents "
        "one employee and each column describes an employee characteristic, "
        "employment condition or outcome."
    )

    structure_df = pd.DataFrame(DATASET_GROUPS)

    add_dataframe_table(
        document,
        structure_df,
        columns=[
            "category",
            "variables",
            "description"
        ],
        column_labels=[
            "Variable Group",
            "Variables",
            "Purpose"
        ]
    )

    # Data preparation
    document.add_heading(
        "3. Data Preparation",
        level=1
    )

    preparation_points = [
        "Converted column names into consistent snake_case format.",
        "Removed duplicate records when present.",
        "Removed extra spaces from text values.",
        "Removed columns that contained only one constant value.",
        "Created attrition_flag where Yes = 1 and No = 0.",
        "Created overtime_flag where Yes = 1 and No = 0."
    ]

    for point in preparation_points:
        paragraph = document.add_paragraph(
            style="List Bullet"
        )

        paragraph.add_run(point)

    constant_columns = quality_report.get(
        "constant_columns_removed",
        []
    )

    document.add_paragraph(
        "Constant columns removed: "
        + (
            ", ".join(constant_columns)
            if constant_columns
            else "None"
        )
    )

    # Key findings
    document.add_heading(
        "4. Key Findings",
        level=1
    )

    for finding in findings:
        paragraph = document.add_paragraph(
            style="List Bullet"
        )

        paragraph.add_run(finding)

    # Department results
    document.add_heading(
        "5. Attrition by Department",
        level=1
    )

    document.add_paragraph(
        "The following table compares department size, employees "
        "who left and the attrition rate within each department."
    )

    department_output = department_df.sort_values(
        "attrition_rate",
        ascending=False
    )

    add_dataframe_table(
        document,
        department_output,
        columns=[
            "department",
            "employees",
            "employees_left",
            "attrition_rate"
        ],
        column_labels=[
            "Department",
            "Employees",
            "Employees Left",
            "Attrition Rate"
        ],
        percentage_columns=[
            "attrition_rate"
        ]
    )

    # Job-role results
    document.add_heading(
        "6. Attrition by Job Role",
        level=1
    )

    document.add_paragraph(
        "Job-role attrition rates provide a more accurate comparison "
        "than departure counts alone because job roles have different "
        "numbers of employees."
    )

    job_role_output = job_role_df.sort_values(
        "attrition_rate",
        ascending=False
    )

    add_dataframe_table(
        document,
        job_role_output,
        columns=[
            "job_role",
            "employees",
            "employees_left",
            "attrition_rate"
        ],
        column_labels=[
            "Job Role",
            "Employees",
            "Employees Left",
            "Attrition Rate"
        ],
        percentage_columns=[
            "attrition_rate"
        ]
    )

    # Charts
    document.add_heading(
        "7. Visual Analysis",
        level=1
    )

    for filename, caption, explanation in CHARTS:
        image_path = FIGURES_DIRECTORY / filename

        if not image_path.exists():
            continue

        document.add_picture(
            str(image_path),
            width=Inches(6.3)
        )

        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        caption_run = caption_paragraph.add_run(
            caption
        )

        caption_run.bold = True

        explanation_paragraph = document.add_paragraph(
            explanation
        )

        explanation_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    # Conclusion
    document.add_heading(
        "8. Conclusion",
        level=1
    )

    document.add_paragraph(
        "The analysis highlights differences in attrition across "
        "departments, job roles, overtime conditions, income and tenure. "
        "These findings can help HR teams identify groups requiring "
        "further investigation and design more targeted employee-retention "
        "strategies."
    )

    document.add_heading(
        "9. Limitations",
        level=1
    )

    limitations = [
        (
            "The dataset identifies statistical relationships but does "
            "not establish direct causation."
        ),
        (
            "The dataset does not include employee comments, exit-interview "
            "responses or detailed organisational policies."
        ),
        (
            "Some categorical ratings are stored as numerical codes, while "
            "their full label definitions are not included in the CSV."
        ),
        (
            "Attrition is imbalanced because substantially more employees "
            "stayed than left. This must be considered during modelling."
        )
    ]

    for limitation in limitations:
        paragraph = document.add_paragraph(
            style="List Bullet"
        )

        paragraph.add_run(limitation)

    document.save(WORD_REPORT)


# --------------------------------------------------
# HTML report
# --------------------------------------------------

def create_html_report(
    cleaned_df: pd.DataFrame,
    eda_summary: dict,
    quality_report: dict,
    department_df: pd.DataFrame,
    job_role_df: pd.DataFrame,
    findings: list[str]
) -> None:
    """Generate a browser-friendly formatted HTML report."""

    structure_df = pd.DataFrame(DATASET_GROUPS)

    department_output = (
        department_df.sort_values(
            "attrition_rate",
            ascending=False
        )
        .copy()
    )

    department_output["attrition_rate"] = (
        department_output["attrition_rate"]
        .map(lambda value: f"{value:.2f}%")
    )

    job_role_output = (
        job_role_df.sort_values(
            "attrition_rate",
            ascending=False
        )
        .copy()
    )

    job_role_output["attrition_rate"] = (
        job_role_output["attrition_rate"]
        .map(lambda value: f"{value:.2f}%")
    )

    findings_html = "\n".join(
        f"<li>{escape(finding)}</li>"
        for finding in findings
    )

    removed_columns = quality_report.get(
        "constant_columns_removed",
        []
    )

    removed_columns_text = (
        ", ".join(removed_columns)
        if removed_columns
        else "None"
    )

    charts_html_parts = []

    for filename, caption, explanation in CHARTS:
        image_path = FIGURES_DIRECTORY / filename

        if not image_path.exists():
            continue

        charts_html_parts.append(
            f"""
            <figure>
                <img
                    src="../figures/{escape(filename)}"
                    alt="{escape(caption)}"
                >
                <figcaption>
                    <strong>{escape(caption)}</strong><br>
                    {escape(explanation)}
                </figcaption>
            </figure>
            """
        )

    charts_html = "\n".join(charts_html_parts)

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta
        name="viewport"
        content="width=device-width, initial-scale=1.0"
    >
    <title>HR Employee Attrition Analysis</title>

    <style>
        body {{
            margin: 0;
            background: #f4f7fa;
            color: #243447;
            font-family: Arial, Helvetica, sans-serif;
            line-height: 1.6;
        }}

        .report {{
            width: min(1100px, calc(100% - 40px));
            margin: 30px auto;
            background: white;
            padding: 48px;
            box-sizing: border-box;
            border-radius: 12px;
            box-shadow: 0 6px 24px rgba(0, 0, 0, 0.08);
        }}

        .title-section {{
            text-align: center;
            padding-bottom: 28px;
            border-bottom: 3px solid #1f4e79;
        }}

        h1 {{
            color: #1f4e79;
            font-size: 34px;
            margin-bottom: 8px;
        }}

        h2 {{
            color: #1f4e79;
            margin-top: 42px;
            padding-bottom: 8px;
            border-bottom: 2px solid #d9eaf7;
        }}

        h3 {{
            color: #2e75b6;
        }}

        .subtitle {{
            color: #526777;
            font-size: 18px;
        }}

        .metrics {{
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 16px;
            margin-top: 24px;
        }}

        .metric-card {{
            padding: 20px;
            background: #eef5fa;
            border-left: 5px solid #2e75b6;
            border-radius: 8px;
        }}

        .metric-card .label {{
            color: #526777;
            font-size: 14px;
        }}

        .metric-card .value {{
            color: #1f4e79;
            font-size: 25px;
            font-weight: bold;
            margin-top: 5px;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 22px 0;
            font-size: 14px;
        }}

        th {{
            background: #1f4e79;
            color: white;
            text-align: left;
            padding: 11px;
        }}

        td {{
            padding: 10px;
            border-bottom: 1px solid #d9e2e8;
            vertical-align: top;
        }}

        tr:nth-child(even) {{
            background: #f7f9fb;
        }}

        .finding-box {{
            background: #f0f8f4;
            border-left: 5px solid #3b8c63;
            padding: 20px 25px;
            border-radius: 8px;
        }}

        .note {{
            background: #fff8e8;
            border-left: 5px solid #d29b22;
            padding: 16px 20px;
            border-radius: 8px;
        }}

        figure {{
            margin: 35px 0;
            text-align: center;
            page-break-inside: avoid;
        }}

        figure img {{
            width: 100%;
            max-width: 850px;
            height: auto;
            border: 1px solid #dde4ea;
            border-radius: 8px;
        }}

        figcaption {{
            margin-top: 10px;
            color: #526777;
        }}

        footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #d9e2e8;
            text-align: center;
            color: #71808c;
            font-size: 13px;
        }}

        @media max-width: 800px {{
            .report {{
                width: 100%;
                margin: 0;
                padding: 25px;
                border-radius: 0;
            }}

            .metrics {{
                grid-template-columns: 1fr;
            }}
        }}

        @media print {{
            body {{
                background: white;
            }}

            .report {{
                width: 100%;
                margin: 0;
                box-shadow: none;
                border-radius: 0;
            }}

            h2 {{
                page-break-after: avoid;
            }}

            table, figure {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>

<body>
    <main class="report">
        <section class="title-section">
            <h1>HR Employee Attrition Analysis</h1>

            <div class="subtitle">
                Data Cleaning, Exploratory Analysis and Key Findings
            </div>

            <p>
                Generated on
                {date.today().strftime("%d %B %Y")}
            </p>
        </section>

        <h2>1. Executive Summary</h2>

        <p>
            This report examines employee attrition using employee
            demographics, job conditions, compensation, satisfaction,
            experience and organisational tenure.
        </p>

        <section class="metrics">
            <div class="metric-card">
                <div class="label">Total Employees</div>
                <div class="value">
                    {eda_summary["total_employees"]:,}
                </div>
            </div>

            <div class="metric-card">
                <div class="label">Employees Who Left</div>
                <div class="value">
                    {eda_summary["employees_left"]:,}
                </div>
            </div>

            <div class="metric-card">
                <div class="label">Employees Who Stayed</div>
                <div class="value">
                    {eda_summary["employees_stayed"]:,}
                </div>
            </div>

            <div class="metric-card">
                <div class="label">Attrition Rate</div>
                <div class="value">
                    {eda_summary["overall_attrition_rate_percent"]:.2f}%
                </div>
            </div>

            <div class="metric-card">
                <div class="label">Average Age</div>
                <div class="value">
                    {eda_summary["average_age"]:.2f}
                </div>
            </div>

            <div class="metric-card">
                <div class="label">Average Monthly Income</div>
                <div class="value">
                    {eda_summary["average_monthly_income"]:,.2f}
                </div>
            </div>
        </section>

        <h2>2. Dataset Overview</h2>

        <p>
            The cleaned dataset contains
            <strong>{cleaned_df.shape[0]:,} employee records</strong>
            and
            <strong>{cleaned_df.shape[1]} variables</strong>.
            Each row represents one employee.
        </p>

        {structure_df.to_html(
            index=False,
            border=0,
            classes="data-table"
        )}

        <h2>3. Data Preparation</h2>

        <ul>
            <li>Standardised column names using snake_case.</li>
            <li>Checked and removed duplicate records.</li>
            <li>Removed extra spaces from text values.</li>
            <li>Removed constant-value columns.</li>
            <li>Created numerical attrition and overtime flags.</li>
        </ul>

        <div class="note">
            <strong>Constant columns removed:</strong>
            {escape(removed_columns_text)}
        </div>

        <h2>4. Key Findings</h2>

        <div class="finding-box">
            <ul>
                {findings_html}
            </ul>
        </div>

        <h2>5. Attrition by Department</h2>

        {department_output.to_html(
            index=False,
            border=0,
            classes="data-table"
        )}

        <h2>6. Attrition by Job Role</h2>

        {job_role_output.to_html(
            index=False,
            border=0,
            classes="data-table"
        )}

        <h2>7. Visual Analysis</h2>

        {charts_html}

        <h2>8. Conclusion</h2>

        <p>
            The analysis identifies meaningful differences in attrition
            across departments, job roles, overtime conditions, employee
            income and organisational tenure. These findings can support
            further investigation and more targeted retention strategies.
        </p>

        <h2>9. Limitations</h2>

        <ul>
            <li>
                The results show statistical relationships and do not
                establish direct causation.
            </li>

            <li>
                The data does not contain exit-interview comments or
                qualitative employee feedback.
            </li>

            <li>
                Some numerical category definitions are not included
                in the original CSV.
            </li>

            <li>
                The target variable is imbalanced because more employees
                stayed than left.
            </li>
        </ul>

        <footer>
            HR Employee Attrition Analysis Report
        </footer>
    </main>
</body>
</html>
"""

    HTML_REPORT.write_text(
        html_content,
        encoding="utf-8"
    )


# --------------------------------------------------
# Main
# --------------------------------------------------

def main() -> None:
    FINAL_DIRECTORY.mkdir(
        parents=True,
        exist_ok=True
    )

    validate_files()

    (
        cleaned_df,
        eda_summary,
        quality_report,
        department_df,
        job_role_df
    ) = load_results()

    findings = build_key_findings(
        cleaned_df,
        eda_summary,
        department_df,
        job_role_df
    )

    create_word_report(
        cleaned_df,
        eda_summary,
        quality_report,
        department_df,
        job_role_df,
        findings
    )

    create_html_report(
        cleaned_df,
        eda_summary,
        quality_report,
        department_df,
        job_role_df,
        findings
    )

    print("\nReports generated successfully.")
    print(f"Word report: {WORD_REPORT}")
    print(f"HTML report: {HTML_REPORT}")


if __name__ == "__main__":
    main()
